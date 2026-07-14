# import os, io, re, uuid
# from datetime import datetime
# from typing import Optional, List

# from fastapi import FastAPI, UploadFile, File, HTTPException, Header, Request
# from fastapi.middleware.cors import CORSMiddleware
# from motor.motor_asyncio import AsyncIOMotorClient
# from bson import ObjectId
# import pdfplumber
# import spacy
# from docx import Document

# MONGO_URL = os.getenv("MONGO_URL", "mongodb://localhost:27017")

# UPLOAD_DIR = "uploads"
# os.makedirs(UPLOAD_DIR, exist_ok=True)

# SKILLS_DICT = {
#     "python","javascript","typescript","java","go","rust","c++","c#","ruby","php","swift","kotlin","scala",
#     "react","vue","angular","next.js","nuxt","svelte","fastapi","django","flask","spring","express","node.js",
#     "react query","redux","tailwind","bootstrap","graphql","rest","grpc",
#     "sql","postgresql","mysql","mongodb","redis","elasticsearch","pandas","numpy","scikit-learn","tensorflow",
#     "pytorch","keras","spark","kafka","airflow","dbt","powerbi","tableau",
#     "docker","kubernetes","terraform","ansible","aws","gcp","azure","ci/cd","github actions","jenkins",
#     "linux","bash","nginx","prometheus","grafana","git","microservices","tdd","system design","api design","nosql",
# }

# app = FastAPI(title="Resume Parser", version="1.0.0")

# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# mongo_client = AsyncIOMotorClient(MONGO_URL)
# db = mongo_client["jat"]

# try:
#     nlp = spacy.load("en_core_web_sm")
# except OSError:
#     from spacy.cli import download
#     download("en_core_web_sm")
#     nlp = spacy.load("en_core_web_sm")


# def extract_text_from_pdf(file_bytes: bytes) -> str:
#     text = ""
#     with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
#         for page in pdf.pages:
#             text += (page.extract_text() or "") + "\n"
#     return text


# def extract_text_from_docx(file_bytes: bytes) -> str:
#     doc = Document(io.BytesIO(file_bytes))
#     return "\n".join(p.text for p in doc.paragraphs)


# def extract_skills(text: str) -> List[str]:
#     lower = text.lower()
#     found = set()

#     for skill in SKILLS_DICT:
#         if re.search(r'\b' + re.escape(skill) + r'\b', lower):
#             found.add(skill)

#     return sorted(found)


# def extract_email(text: str) -> Optional[str]:
#     m = re.search(r'[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}', text)
#     return m.group(0) if m else None


# def extract_phone(text: str) -> Optional[str]:
#     m = re.search(r'[\+]?[\d\s\-().]{7,15}', text)
#     return m.group(0).strip() if m else None


# def extract_name(text: str) -> Optional[str]:
#     for line in text.split('\n'):
#         line = line.strip()

#         if line and len(line.split()) >= 2:
#             return line[:60]

#     return None


# def parse_resume(text: str) -> dict:
#     return {
#         "skills": extract_skills(text),
#         "name": extract_name(text),
#         "email": extract_email(text),
#         "phone": extract_phone(text),
#     }


# @app.post("/resumes/upload")
# async def upload_resume(
#     file: UploadFile = File(...),
#     x_user_id: str = Header(...)
# ):
#     file_bytes = await file.read()

#     filename = file.filename.lower()

#     if filename.endswith(".pdf"):
#         raw_text = extract_text_from_pdf(file_bytes)

#     elif filename.endswith(".docx"):
#         raw_text = extract_text_from_docx(file_bytes)

#     else:
#         raise HTTPException(
#             status_code=400,
#             detail="Only PDF and DOCX supported"
#         )

#     unique_name = f"{uuid.uuid4()}_{file.filename}"

#     file_path = os.path.join(UPLOAD_DIR, unique_name)

#     with open(file_path, "wb") as f:
#         f.write(file_bytes)

#     parsed = parse_resume(raw_text)

#     doc = {
#         "user_id": x_user_id,
#         "file_url": file_path,
#         "original_filename": file.filename,
#         "raw_text": raw_text,
#         "parsed": parsed,
#         "created_at": datetime.utcnow(),
#     }

#     result = await db.resumes.insert_one(doc)

#     return {
#         "id": str(result.inserted_id),
#         "user_id": x_user_id,
#         "file_url": file_path,
#         "original_filename": file.filename,
#         "parsed": parsed,
#         "created_at": doc["created_at"].isoformat(),
#     }


# @app.get("/resumes")
# async def list_resumes(x_user_id: str = Header(...)):
#     cursor = db.resumes.find(
#         {"user_id": x_user_id}
#     ).sort("created_at", -1)

#     resumes = []

#     async for doc in cursor:
#         doc["id"] = str(doc.pop("_id"))
#         doc["created_at"] = doc["created_at"].isoformat()
#         doc.pop("raw_text", None)

#         resumes.append(doc)

#     return resumes

# @app.delete("/resumes/{resume_id}")
# async def delete_resume(
#     resume_id: str,
#     x_user_id: str = Header(...)
# ):
#     result = await db.resumes.delete_one({
#         "_id": ObjectId(resume_id),
#         "user_id": x_user_id
#     })

#     if result.deleted_count == 0:
#         raise HTTPException(
#             status_code=404,
#             detail="Resume not found"
#         )

#     return {
#         "message": "Resume deleted successfully"
#     }


# @app.get("/health")
# async def health():
#     return {
#         "status": "ok",
#         "service": "resume_parser"
#     }



# import io
# import os
# import re
# import uuid
# from datetime import datetime, timezone
# from typing import List, Optional

# import pdfplumber
# from bson import ObjectId
# from bson.errors import InvalidId
# from docx import Document
# from fastapi import (
#     FastAPI,
#     File,
#     Header,
#     HTTPException,
#     UploadFile,
# )
# from fastapi.middleware.cors import CORSMiddleware
# from fastapi.responses import FileResponse
# from motor.motor_asyncio import AsyncIOMotorClient


# # ============================================================
# # CONFIGURATION
# # ============================================================

# MONGO_URL = os.getenv(
#     "MONGO_URL",
#     "mongodb://localhost:27017",
# )

# BASE_DIR = os.path.dirname(
#     os.path.abspath(__file__)
# )

# UPLOAD_DIR = os.path.join(
#     BASE_DIR,
#     "uploads",
# )

# os.makedirs(
#     UPLOAD_DIR,
#     exist_ok=True,
# )


# # ============================================================
# # SKILLS DICTIONARY
# # ============================================================

# SKILLS_DICT = {
#     "python",
#     "javascript",
#     "typescript",
#     "java",
#     "go",
#     "rust",
#     "c++",
#     "c#",
#     "ruby",
#     "php",
#     "swift",
#     "kotlin",
#     "scala",

#     "react",
#     "vue",
#     "angular",
#     "next.js",
#     "nuxt",
#     "svelte",

#     "fastapi",
#     "django",
#     "flask",
#     "spring",
#     "spring boot",
#     "express",
#     "node.js",

#     "react query",
#     "redux",
#     "tailwind",
#     "bootstrap",

#     "graphql",
#     "rest",
#     "rest api",
#     "grpc",

#     "sql",
#     "postgresql",
#     "mysql",
#     "mongodb",
#     "redis",
#     "elasticsearch",

#     "pandas",
#     "numpy",
#     "scikit-learn",
#     "tensorflow",
#     "pytorch",
#     "keras",

#     "spark",
#     "kafka",
#     "airflow",
#     "dbt",

#     "power bi",
#     "powerbi",
#     "tableau",

#     "docker",
#     "kubernetes",
#     "terraform",
#     "ansible",

#     "aws",
#     "gcp",
#     "azure",

#     "ci/cd",
#     "github actions",
#     "jenkins",

#     "linux",
#     "bash",
#     "nginx",

#     "prometheus",
#     "grafana",

#     "git",
#     "github",

#     "microservices",
#     "tdd",
#     "system design",
#     "api design",
#     "nosql",
# }


# # ============================================================
# # FASTAPI APPLICATION
# # ============================================================

# app = FastAPI(
#     title="Resume Parser",
#     version="1.2.0",
# )


# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],
#     allow_methods=["*"],
#     allow_headers=["*"],
# )


# # ============================================================
# # DATABASE
# # ============================================================

# mongo_client = AsyncIOMotorClient(
#     MONGO_URL
# )

# db = mongo_client["jat"]


# # ============================================================
# # OBJECT ID HELPER
# # ============================================================

# def get_object_id(
#     value: str,
# ) -> ObjectId:

#     try:

#         return ObjectId(value)

#     except (
#         InvalidId,
#         TypeError,
#     ):

#         raise HTTPException(
#             status_code=400,
#             detail="Invalid resume ID",
#         )


# # ============================================================
# # RESUME OWNERSHIP HELPER
# # ============================================================

# async def get_resume_document(
#     resume_id: str,
#     user_id: str,
# ) -> dict:

#     object_id = get_object_id(
#         resume_id
#     )

#     resume = await db.resumes.find_one(
#         {
#             "_id": object_id,
#             "user_id": user_id,
#         }
#     )

#     if not resume:

#         raise HTTPException(
#             status_code=404,
#             detail="Resume not found",
#         )

#     return resume


# # ============================================================
# # FILE PATH HELPERS
# # ============================================================

# def build_upload_path(
#     stored_filename: str,
# ) -> str:
#     """
#     Convert a stored server filename into a safe absolute path.

#     MongoDB should contain:

#         server_filename:
#         65e7....pdf

#     The frontend never receives this value.
#     """

#     if not stored_filename:

#         raise HTTPException(
#             status_code=404,
#             detail="Resume filename not found",
#         )


#     safe_filename = os.path.basename(
#         stored_filename
#     )


#     file_path = os.path.abspath(
#         os.path.join(
#             UPLOAD_DIR,
#             safe_filename,
#         )
#     )


#     upload_directory = os.path.abspath(
#         UPLOAD_DIR
#     )


#     try:

#         common_path = os.path.commonpath(
#             [
#                 file_path,
#                 upload_directory,
#             ]
#         )

#     except ValueError:

#         raise HTTPException(
#             status_code=403,
#             detail="Invalid resume file path",
#         )


#     if common_path != upload_directory:

#         raise HTTPException(
#             status_code=403,
#             detail="Invalid resume file path",
#         )


#     return file_path


# def resolve_resume_file(
#     resume: dict,
# ) -> str:
#     """
#     Resolve physical file location.

#     Supports:

#     NEW RECORDS:
#         server_filename

#     OLD RECORDS:
#         file_path
#         file_url

#     This allows migration without breaking
#     already-uploaded resumes.
#     """

#     # --------------------------------------------------------
#     # NEW STORAGE FORMAT
#     # --------------------------------------------------------

#     server_filename = resume.get(
#         "server_filename"
#     )


#     if server_filename:

#         return build_upload_path(
#             server_filename
#         )


#     # --------------------------------------------------------
#     # OLD STORAGE FORMAT
#     # --------------------------------------------------------

#     stored_path = (
#         resume.get("file_path")
#         or resume.get("file_url")
#     )


#     if not stored_path:

#         raise HTTPException(
#             status_code=404,
#             detail="Resume file location not found",
#         )


#     # Absolute path from old records.

#     if os.path.isabs(stored_path):

#         candidate = os.path.abspath(
#             stored_path
#         )

#     else:

#         # Example:
#         # uploads/file.pdf

#         candidate = os.path.abspath(
#             os.path.join(
#                 BASE_DIR,
#                 stored_path,
#             )
#         )


#     # If old path no longer works,
#     # try only its filename inside UPLOAD_DIR.

#     if not os.path.isfile(candidate):

#         candidate = build_upload_path(
#             os.path.basename(
#                 stored_path
#             )
#         )


#     upload_directory = os.path.abspath(
#         UPLOAD_DIR
#     )


#     try:

#         common_path = os.path.commonpath(
#             [
#                 candidate,
#                 upload_directory,
#             ]
#         )

#     except ValueError:

#         raise HTTPException(
#             status_code=403,
#             detail="Invalid resume file path",
#         )


#     if common_path != upload_directory:

#         raise HTTPException(
#             status_code=403,
#             detail="Invalid resume file path",
#         )


#     return candidate


# # ============================================================
# # PDF TEXT EXTRACTION
# # ============================================================

# def extract_text_from_pdf(
#     file_bytes: bytes,
# ) -> str:

#     text_parts = []


#     try:

#         with pdfplumber.open(
#             io.BytesIO(file_bytes)
#         ) as pdf:

#             for page in pdf.pages:

#                 page_text = (
#                     page.extract_text()
#                 )


#                 if page_text:

#                     text_parts.append(
#                         page_text
#                     )


#     except Exception:

#         raise HTTPException(
#             status_code=400,
#             detail="Unable to read PDF file",
#         )


#     return "\n".join(
#         text_parts
#     )


# # ============================================================
# # DOCX TEXT EXTRACTION
# # ============================================================

# def extract_text_from_docx(
#     file_bytes: bytes,
# ) -> str:

#     try:

#         document = Document(
#             io.BytesIO(file_bytes)
#         )


#         return "\n".join(
#             paragraph.text
#             for paragraph
#             in document.paragraphs
#             if paragraph.text.strip()
#         )


#     except Exception:

#         raise HTTPException(
#             status_code=400,
#             detail="Unable to read DOCX file",
#         )


# # ============================================================
# # SKILL EXTRACTION
# # ============================================================

# def extract_skills(
#     text: str,
# ) -> List[str]:

#     normalized_text = (
#         text
#         .lower()
#         .replace("–", "-")
#         .replace("—", "-")
#     )


#     found = set()


#     for skill in SKILLS_DICT:

#         pattern = (
#             r"(?<![a-zA-Z0-9])"
#             + re.escape(skill)
#             + r"(?![a-zA-Z0-9])"
#         )


#         if re.search(
#             pattern,
#             normalized_text,
#         ):

#             found.add(skill)


#     return sorted(found)


# # ============================================================
# # EMAIL EXTRACTION
# # ============================================================

# def extract_email(
#     text: str,
# ) -> Optional[str]:

#     match = re.search(
#         r"[\w.+-]+@[\w-]+(?:\.[\w-]+)+",
#         text,
#     )


#     return (
#         match.group(0)
#         if match
#         else None
#     )


# # ============================================================
# # PHONE EXTRACTION
# # ============================================================

# def extract_phone(
#     text: str,
# ) -> Optional[str]:

#     matches = re.findall(
#         r"(?:\+?\d[\d\s\-().]{7,18}\d)",
#         text,
#     )


#     for value in matches:

#         digits = re.sub(
#             r"\D",
#             "",
#             value,
#         )


#         if 10 <= len(digits) <= 15:

#             return value.strip()


#     return None


# # ============================================================
# # NAME EXTRACTION
# # ============================================================

# def extract_name(
#     text: str,
# ) -> Optional[str]:

#     for line in text.splitlines():

#         line = line.strip()


#         if not line:

#             continue


#         if "@" in line:

#             continue


#         if re.search(
#             r"\d{5,}",
#             line,
#         ):

#             continue


#         words = line.split()


#         if 2 <= len(words) <= 6:

#             return line[:60]


#     return None


# # ============================================================
# # PARSE RESUME
# # ============================================================

# def parse_resume(
#     text: str,
# ) -> dict:

#     return {
#         "skills": extract_skills(
#             text
#         ),

#         "name": extract_name(
#             text
#         ),

#         "email": extract_email(
#             text
#         ),

#         "phone": extract_phone(
#             text
#         ),
#     }


# # ============================================================
# # UPLOAD RESUME
# # ============================================================

# @app.post("/resumes/upload")
# async def upload_resume(
#     file: UploadFile = File(...),
#     x_user_id: str = Header(...),
# ):

#     if not file.filename:

#         raise HTTPException(
#             status_code=400,
#             detail="Filename is missing",
#         )


#     extension = os.path.splitext(
#         file.filename
#     )[1].lower()


#     if extension not in {
#         ".pdf",
#         ".docx",
#     }:

#         raise HTTPException(
#             status_code=400,
#             detail="Only PDF and DOCX files are supported",
#         )


#     file_bytes = await file.read()


#     if not file_bytes:

#         raise HTTPException(
#             status_code=400,
#             detail="Uploaded file is empty",
#         )


#     # --------------------------------------------------------
#     # EXTRACT TEXT
#     # --------------------------------------------------------

#     if extension == ".pdf":

#         raw_text = extract_text_from_pdf(
#             file_bytes
#         )

#     else:

#         raw_text = extract_text_from_docx(
#             file_bytes
#         )


#     # --------------------------------------------------------
#     # PARSE RESUME
#     # --------------------------------------------------------

#     parsed = parse_resume(
#         raw_text
#     )


#     # --------------------------------------------------------
#     # CREATE SAFE SERVER FILENAME
#     # --------------------------------------------------------

#     server_filename = (
#         f"{uuid.uuid4().hex}"
#         f"{extension}"
#     )


#     file_path = build_upload_path(
#         server_filename
#     )


#     # --------------------------------------------------------
#     # SAVE PHYSICAL FILE
#     # --------------------------------------------------------

#     try:

#         with open(
#             file_path,
#             "wb",
#         ) as saved_file:

#             saved_file.write(
#                 file_bytes
#             )


#     except OSError:

#         raise HTTPException(
#             status_code=500,
#             detail="Unable to save resume file",
#         )


#     now = datetime.now(
#         timezone.utc
#     )


#     # --------------------------------------------------------
#     # MONGODB DOCUMENT
#     # --------------------------------------------------------

#     document = {
#         "user_id":
#             x_user_id,

#         "server_filename":
#             server_filename,

#         "original_filename":
#             file.filename,

#         "content_type":
#             (
#                 file.content_type
#                 or (
#                     "application/pdf"
#                     if extension == ".pdf"
#                     else
#                     "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#                 )
#             ),

#         "raw_text":
#             raw_text,

#         "parsed":
#             parsed,

#         "created_at":
#             now,

#         "updated_at":
#             now,
#     }


#     # --------------------------------------------------------
#     # INSERT MONGODB DOCUMENT
#     # --------------------------------------------------------

#     try:

#         result = await db.resumes.insert_one(
#             document
#         )


#     except Exception:

#         # Avoid orphan files.

#         if os.path.isfile(
#             file_path
#         ):

#             try:

#                 os.remove(
#                     file_path
#                 )

#             except OSError:

#                 pass


#         raise HTTPException(
#             status_code=500,
#             detail="Unable to save resume information",
#         )


#     # --------------------------------------------------------
#     # RESPONSE
#     # --------------------------------------------------------

#     return {
#         "id":
#             str(result.inserted_id),

#         "user_id":
#             x_user_id,

#         "original_filename":
#             file.filename,

#         "parsed":
#             parsed,

#         "created_at":
#             now.isoformat(),
#     }


# # ============================================================
# # LIST RESUMES
# # ============================================================

# @app.get("/resumes")
# async def list_resumes(
#     x_user_id: str = Header(...),
# ):

#     cursor = (
#         db.resumes
#         .find(
#             {
#                 "user_id":
#                     x_user_id,
#             }
#         )
#         .sort(
#             "created_at",
#             -1,
#         )
#     )


#     resumes = []


#     async for document in cursor:

#         created_at = document.get(
#             "created_at"
#         )


#         resumes.append(
#             {
#                 "id":
#                     str(
#                         document["_id"]
#                     ),

#                 "original_filename":
#                     document.get(
#                         "original_filename",
#                         "Resume",
#                     ),

#                 "parsed":
#                     document.get(
#                         "parsed",
#                         {},
#                     ),

#                 "created_at":
#                     (
#                         created_at.isoformat()
#                         if created_at
#                         else None
#                     ),
#             }
#         )


#     return resumes


# # ============================================================
# # PREVIEW RESUME
# # ============================================================

# @app.get("/resumes/{resume_id}/preview")
# async def preview_resume(
#     resume_id: str,
#     x_user_id: str = Header(...),
# ):

#     resume = await get_resume_document(
#         resume_id,
#         x_user_id,
#     )


#     file_path = resolve_resume_file(
#         resume
#     )


#     if not os.path.isfile(
#         file_path
#     ):

#         raise HTTPException(
#             status_code=404,
#             detail="Resume file not found on server",
#         )


#     filename = resume.get(
#         "original_filename",
#         "resume",
#     )


#     content_type = resume.get(
#         "content_type"
#     )


#     if not content_type:

#         if filename.lower().endswith(
#             ".pdf"
#         ):

#             content_type = (
#                 "application/pdf"
#             )

#         else:

#             content_type = (
#                 "application/vnd."
#                 "openxmlformats-officedocument."
#                 "wordprocessingml.document"
#             )


#     return FileResponse(
#         path=file_path,

#         media_type=
#             content_type,

#         filename=
#             filename,

#         content_disposition_type=
#             "inline",
#     )


# # ============================================================
# # DOWNLOAD RESUME
# # ============================================================

# @app.get("/resumes/{resume_id}/download")
# async def download_resume(
#     resume_id: str,
#     x_user_id: str = Header(...),
# ):

#     resume = await get_resume_document(
#         resume_id,
#         x_user_id,
#     )


#     file_path = resolve_resume_file(
#         resume
#     )


#     if not os.path.isfile(
#         file_path
#     ):

#         raise HTTPException(
#             status_code=404,
#             detail="Resume file not found on server",
#         )


#     filename = resume.get(
#         "original_filename",
#         "resume",
#     )


#     content_type = (
#         resume.get(
#             "content_type"
#         )
#         or
#         "application/octet-stream"
#     )


#     return FileResponse(
#         path=file_path,

#         media_type=
#             content_type,

#         filename=
#             filename,

#         content_disposition_type=
#             "attachment",
#     )


# # ============================================================
# # DELETE RESUME
# # ============================================================

# @app.delete("/resumes/{resume_id}")
# async def delete_resume(
#     resume_id: str,
#     x_user_id: str = Header(...),
# ):

#     resume = await get_resume_document(
#         resume_id,
#         x_user_id,
#     )


#     # --------------------------------------------------------
#     # RESOLVE FILE BEFORE DELETING DATABASE RECORD
#     # --------------------------------------------------------

#     file_path = None


#     try:

#         file_path = resolve_resume_file(
#             resume
#         )

#     except HTTPException:

#         # Still allow stale MongoDB records
#         # to be deleted.

#         file_path = None


#     object_id = get_object_id(
#         resume_id
#     )


#     # --------------------------------------------------------
#     # DELETE MONGODB RECORD
#     # --------------------------------------------------------

#     result = await db.resumes.delete_one(
#         {
#             "_id":
#                 object_id,

#             "user_id":
#                 x_user_id,
#         }
#     )


#     if result.deleted_count == 0:

#         raise HTTPException(
#             status_code=404,
#             detail="Resume not found",
#         )


#     # --------------------------------------------------------
#     # DELETE PHYSICAL FILE
#     # --------------------------------------------------------

#     if (
#         file_path
#         and os.path.isfile(
#             file_path
#         )
#     ):

#         try:

#             os.remove(
#                 file_path
#             )

#         except OSError:

#             # In production this should be logged.
#             pass


#     return {
#         "message":
#             "Resume deleted successfully",
#     }


# # ============================================================
# # HEALTH
# # ============================================================

# @app.get("/health")
# async def health():

#     return {
#         "status":
#             "ok",

#         "service":
#             "resume_parser",
#     }




import os
import io
import re
import uuid
from datetime import datetime, timezone
from typing import List

import pdfplumber

from fastapi import (
    FastAPI,
    UploadFile,
    File,
    HTTPException,
    Header,
)

from fastapi.middleware.cors import CORSMiddleware

from fastapi.responses import StreamingResponse

from motor.motor_asyncio import AsyncIOMotorClient

from bson import ObjectId

from docx import Document

# from minio import Minio
# from minio.error import S3Error

import httpx


# ============================================================
# ENVIRONMENT CONFIGURATION
# ============================================================

MONGO_URL = os.getenv(
    "MONGO_URL",
    "mongodb+srv://jat_db_user:Pk%407975463006@cluster0.6hy7bvy.mongodb.net/?appName=Cluster0",
)

# MINIO_ENDPOINT = os.getenv(
#     "MINIO_ENDPOINT",
#     # "localhost:9000",
#     "minio:9000",
# )

# MINIO_ACCESS_KEY = os.getenv(
#     "MINIO_ACCESS_KEY",
#     "minioadmin",
# )

# MINIO_SECRET_KEY = os.getenv(
#     "MINIO_SECRET_KEY",
#     "minioadmin",
# )

# MINIO_BUCKET = os.getenv(
#     "MINIO_BUCKET",
#     "resumes",
# )

# MINIO_SECURE = (
#     os.getenv("MINIO_SECURE", "false").lower()
#     == "true"
# )


# SUPABASE_URL = os.getenv("SUPABASE_URL", "").rstrip("/")

# SUPABASE_SERVICE_KEY = os.getenv(
#     "SUPABASE_SERVICE_KEY",
#     "",
# )

# SUPABASE_BUCKET = os.getenv(
#     "SUPABASE_BUCKET",
#     "resumes",
# )


SUPABASE_URL = os.getenv(
    "SUPABASE_URL",
    "",
).strip().rstrip("/")

SUPABASE_SERVICE_KEY = os.getenv(
    "SUPABASE_SERVICE_KEY",
    "",
).strip()

SUPABASE_BUCKET = os.getenv(
    "SUPABASE_BUCKET",
    "resumes",
).strip().strip("/")


if not SUPABASE_URL:
    raise RuntimeError(
        "SUPABASE_URL environment variable is missing"
    )


if not SUPABASE_SERVICE_KEY:
    raise RuntimeError(
        "SUPABASE_SERVICE_KEY environment variable is missing"
    )


# ============================================================
# FASTAPI
# ============================================================

app = FastAPI(
    title="Resume Parser",
    version="2.0.0",
)


app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ============================================================
# MONGODB
# ============================================================

mongo_client = AsyncIOMotorClient(MONGO_URL)

db = mongo_client["jat"]


# ============================================================
# MINIO
# ============================================================

# minio_client = Minio(
#     MINIO_ENDPOINT,
#     access_key=MINIO_ACCESS_KEY,
#     secret_key=MINIO_SECRET_KEY,
#     secure=MINIO_SECURE,
# )


# ============================================================
# SKILLS
# ============================================================

SKILLS_DICT = {

    # PROGRAMMING

    "python",
    "javascript",
    "typescript",
    "java",
    "go",
    "rust",
    "c++",
    "c#",
    "ruby",
    "php",
    "swift",
    "kotlin",
    "scala",

    # FRONTEND

    "html",
    "css",
    "react",
    "react.js",
    "reactjs",
    "vue",
    "angular",
    "next.js",
    "svelte",
    "bootstrap",
    "tailwind",

    # BACKEND

    "fastapi",
    "django",
    "flask",
    "spring",
    "spring boot",
    "express",
    "node.js",
    "nodejs",

    # DATABASE

    "sql",
    "mysql",
    "postgresql",
    "mongodb",
    "redis",
    "sqlite",

    # DATA

    "pandas",
    "numpy",
    "scikit-learn",
    "tensorflow",
    "pytorch",
    "keras",
    "power bi",
    "powerbi",
    "tableau",

    # DEVOPS

    "docker",
    "kubernetes",
    "terraform",
    "ansible",
    "aws",
    "azure",
    "gcp",
    "jenkins",
    "github actions",
    "ci/cd",

    # TOOLS

    "git",
    "github",
    "postman",
    "linux",
    "bash",

    # ARCHITECTURE

    "microservices",
    "rest",
    "rest api",
    "graphql",
    "grpc",
    "system design",
    "api design",
    "tdd",
}


# ============================================================
# NORMALIZATION
# ============================================================

SKILL_ALIASES = {

    "react.js": "react",
    "reactjs": "react",

    "nodejs": "node.js",

    "spring": "spring boot",

    "powerbi": "power bi",

    "rest": "rest api",
}


def normalize_skill(skill: str) -> str:

    skill = skill.strip().lower()

    return SKILL_ALIASES.get(
        skill,
        skill,
    )


# ============================================================
# EXTRACT SKILLS
# ============================================================

def extract_skills(text: str) -> List[str]:

    if not text:
        return []

    lower_text = text.lower()

    skills = set()

    for skill in SKILLS_DICT:

        pattern = (
            r"(?<!\w)"
            + re.escape(skill)
            + r"(?!\w)"
        )

        if re.search(
            pattern,
            lower_text,
            re.IGNORECASE,
        ):

            skills.add(
                normalize_skill(skill)
            )

    return sorted(skills)


# ============================================================
# PDF PARSER
# ============================================================

def parse_pdf(file_bytes: bytes) -> str:

    text_parts = []

    try:

        with pdfplumber.open(
            io.BytesIO(file_bytes)
        ) as pdf:

            for page in pdf.pages:

                page_text = page.extract_text()

                if page_text:
                    text_parts.append(page_text)

    except Exception as exc:

        raise HTTPException(
            status_code=400,
            detail=f"PDF parsing failed: {exc}",
        )

    return "\n".join(text_parts)


# ============================================================
# DOCX PARSER
# ============================================================

def parse_docx(file_bytes: bytes) -> str:

    try:

        document = Document(
            io.BytesIO(file_bytes)
        )

        text_parts = [
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]

        # Extract text from tables too

        for table in document.tables:

            for row in table.rows:

                for cell in row.cells:

                    value = cell.text.strip()

                    if value:
                        text_parts.append(value)

        return "\n".join(text_parts)

    except Exception as exc:

        raise HTTPException(
            status_code=400,
            detail=f"DOCX parsing failed: {exc}",
        )


# ============================================================
# TEXT PARSER
# ============================================================

def parse_txt(file_bytes: bytes) -> str:

    try:

        return file_bytes.decode(
            "utf-8",
            errors="ignore",
        )

    except Exception as exc:

        raise HTTPException(
            status_code=400,
            detail=f"TXT parsing failed: {exc}",
        )


# ============================================================
# FILE PARSER
# ============================================================

def parse_resume(
    filename: str,
    file_bytes: bytes,
) -> str:

    extension = os.path.splitext(
        filename
    )[1].lower()

    if extension == ".pdf":

        return parse_pdf(file_bytes)

    if extension == ".docx":

        return parse_docx(file_bytes)

    if extension == ".txt":

        return parse_txt(file_bytes)

    raise HTTPException(
        status_code=400,
        detail=(
            "Unsupported file type. "
            "Only PDF, DOCX and TXT are supported."
        ),
    )


# ============================================================
# MONGODB SERIALIZER
# ============================================================

def serialize_resume(document):

    return {

        "id":
            str(document["_id"]),

        "original_filename":
            document.get(
                "original_filename",
                "resume",
            ),

        "content_type":
            document.get(
                "content_type",
                "application/octet-stream",
            ),

        "file_size":
            document.get(
                "file_size",
                0,
            ),

        "parsed":
            document.get(
                "parsed",
                {},
            ),

        "created_at":
            document.get(
                "created_at"
            ),

        "updated_at":
            document.get(
                "updated_at"
            ),
    }


# ============================================================
# STARTUP
# ============================================================

# @app.on_event("startup")
# async def startup():

#     try:

#         if not minio_client.bucket_exists(
#             MINIO_BUCKET
#         ):

#             minio_client.make_bucket(
#                 MINIO_BUCKET
#             )

#     except Exception as exc:

#         print(
#             "MINIO STARTUP ERROR:",
#             exc,
#         )



def get_supabase_headers(
    content_type: str = None,
):

    headers = {
        "Authorization":
            f"Bearer {SUPABASE_SERVICE_KEY}",

        "apikey":
            SUPABASE_SERVICE_KEY,
    }

    if content_type:
        headers["Content-Type"] = content_type

    return headers


# def get_storage_url(
#     object_name: str,
# ) -> str:

#     return (
#         f"{SUPABASE_URL}"
#         f"/storage/v1/object/"
#         f"{SUPABASE_BUCKET}/"
#         f"{object_name}"
#     )


# async def upload_storage_file(
#     object_name: str,
#     file_bytes: bytes,
#     content_type: str,
# ):

#     url = get_storage_url(
#         object_name
#     )

#     async with httpx.AsyncClient(
#         timeout=60.0
#     ) as client:

#         response = await client.post(
#             url,
#             headers=get_supabase_headers(
#                 content_type
#             ),
#             content=file_bytes,
#         )

#     if response.status_code not in (
#         200,
#         201,
#     ):

#         raise HTTPException(
#             status_code=500,
#             detail=(
#                 "File storage failed: "
#                 f"{response.text}"
#             ),
#         )



def get_storage_url(object_name: str) -> str:
    clean_object_name = object_name.lstrip("/")

    return (
        f"{SUPABASE_URL}"
        f"/storage/v1/object/"
        f"{SUPABASE_BUCKET}/"
        f"{clean_object_name}"
    )


async def upload_storage_file(
    object_name: str,
    file_bytes: bytes,
    content_type: str,
):
    url = get_storage_url(object_name)

    headers = {
        "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
        "apikey": SUPABASE_SERVICE_KEY,
        "Content-Type": content_type,
        "x-upsert": "false",
    }

    print("SUPABASE BASE URL:", repr(SUPABASE_URL))
    print("SUPABASE BUCKET:", repr(SUPABASE_BUCKET))
    print("SUPABASE OBJECT:", repr(object_name))
    print("SUPABASE FINAL URL:", repr(url))

    async with httpx.AsyncClient(timeout=60.0) as client:
        response = await client.post(
            url,
            headers=headers,
            content=file_bytes,
        )

    print(
        "SUPABASE RESPONSE:",
        response.status_code,
        response.text,
    )

    if response.status_code not in (200, 201):
        raise HTTPException(
            status_code=500,
            detail=f"File storage failed: {response.text}",
        )
    

async def read_storage_file(
    object_name: str,
) -> bytes:

    url = get_storage_url(
        object_name
    )

    async with httpx.AsyncClient(
        timeout=60.0
    ) as client:

        response = await client.get(
            url,
            headers=get_supabase_headers(),
        )

    if response.status_code != 200:

        raise HTTPException(
            status_code=404,
            detail=(
                "Stored resume file "
                "not found"
            ),
        )

    return response.content


async def delete_storage_file(
    object_name: str,
):

    url = get_storage_url(
        object_name
    )

    async with httpx.AsyncClient(
        timeout=60.0
    ) as client:

        response = await client.delete(
            url,
            headers=get_supabase_headers(),
        )

    if response.status_code not in (
        200,
        204,
    ):

        raise HTTPException(
            status_code=500,
            detail=(
                "Unable to delete "
                "stored file: "
                f"{response.text}"
            ),
        )
    

# ============================================================
# LIST RESUMES
# ============================================================

@app.get("/resumes")
async def list_resumes(
    x_user_id: str = Header(...),
):

    cursor = db.resumes.find(
        {
            "user_id": x_user_id,
        }
    ).sort(
        "created_at",
        -1,
    )

    resumes = []

    async for document in cursor:

        resumes.append(
            serialize_resume(document)
        )

    return resumes


# ============================================================
# UPLOAD RESUME
# ============================================================

@app.post("/resumes/upload")
async def upload_resume(

    file: UploadFile = File(...),

    x_user_id: str = Header(...),

):

    if not file.filename:

        raise HTTPException(
            status_code=400,
            detail="Filename missing",
        )


    extension = os.path.splitext(
        file.filename
    )[1].lower()


    allowed_extensions = {
        ".pdf",
        ".docx",
        ".txt",
    }


    if extension not in allowed_extensions:

        raise HTTPException(
            status_code=400,
            detail=(
                "Only PDF, DOCX and TXT files are supported"
            ),
        )


    file_bytes = await file.read()


    if not file_bytes:

        raise HTTPException(
            status_code=400,
            detail="Uploaded file is empty",
        )


    MAX_FILE_SIZE = 10 * 1024 * 1024


    if len(file_bytes) > MAX_FILE_SIZE:

        raise HTTPException(
            status_code=413,
            detail="Maximum file size is 10 MB",
        )


    # Parse resume

    raw_text = parse_resume(
        file.filename,
        file_bytes,
    )


    if not raw_text.strip():

        raise HTTPException(
            status_code=400,
            detail=(
                "No readable text found in resume"
            ),
        )


    skills = extract_skills(
        raw_text
    )


    # Unique object name prevents filename collisions

    object_name = (

        f"{x_user_id}/"

        f"{uuid.uuid4()}"

        f"{extension}"

    )


    await upload_storage_file(
        object_name=object_name,
        file_bytes=file_bytes,
        content_type=(
            file.content_type
            or "application/octet-stream"
        ),
    )

    now = datetime.now(timezone.utc)


    resume_document = {

        "user_id":
            x_user_id,

        "original_filename":
            file.filename,

        "object_name":
            object_name,

        "content_type":
            file.content_type
            or "application/octet-stream",

        "file_size":
            len(file_bytes),

        "raw_text":
            raw_text,

        "parsed": {

            "skills":
                skills,

        },

        "created_at":
            now,

        "updated_at":
            now,

    }


    result = await db.resumes.insert_one(
        resume_document
    )


    resume_document["_id"] = result.inserted_id


    return serialize_resume(
        resume_document
    )


# ============================================================
# GET RESUME DOCUMENT
# ============================================================

async def get_resume_document(
    resume_id: str,
    user_id: str,
):

    if not ObjectId.is_valid(resume_id):

        raise HTTPException(
            status_code=400,
            detail="Invalid resume ID",
        )


    document = await db.resumes.find_one(

        {

            "_id":
                ObjectId(resume_id),

            "user_id":
                user_id,

        }

    )


    if not document:

        raise HTTPException(
            status_code=404,
            detail="Resume not found",
        )


    return document


# ============================================================
# READ MINIO FILE
# ============================================================

# def read_minio_file(
#     object_name: str,
# ):

#     response = None

#     try:

#         response = minio_client.get_object(

#             MINIO_BUCKET,

#             object_name,

#         )


#         return response.read()


#     except S3Error as exc:

#         raise HTTPException(
#             status_code=404,
#             detail=(
#                 f"Stored resume file not found: {exc}"
#             ),
#         )


#     finally:

#         if response:

#             response.close()

#             response.release_conn()


# ============================================================
# PREVIEW
# ============================================================

@app.get("/resumes/{resume_id}/preview")
async def preview_resume(

    resume_id: str,

    x_user_id: str = Header(...),

):

    document = await get_resume_document(

        resume_id,

        x_user_id,

    )


    # file_bytes = read_minio_file(

    #     document["object_name"]

    # )


    file_bytes = await read_storage_file(
    document["object_name"]
)


    content_type = document.get(

        "content_type",

        "application/octet-stream",

    )


    filename = document.get(

        "original_filename",

        "resume",

    )


    return StreamingResponse(

        io.BytesIO(file_bytes),

        media_type=content_type,

        headers={

            "Content-Disposition":

                f'inline; filename="{filename}"'

        },

    )


# ============================================================
# DOWNLOAD
# ============================================================

@app.get("/resumes/{resume_id}/download")
async def download_resume(

    resume_id: str,

    x_user_id: str = Header(...),

):

    document = await get_resume_document(

        resume_id,

        x_user_id,

    )


    # file_bytes = read_minio_file(

    #     document["object_name"]

    # )

    file_bytes = await read_storage_file(
    document["object_name"]
)


    content_type = document.get(

        "content_type",

        "application/octet-stream",

    )


    filename = document.get(

        "original_filename",

        "resume",

    )


    return StreamingResponse(

        io.BytesIO(file_bytes),

        media_type=content_type,

        headers={

            "Content-Disposition":

                f'attachment; filename="{filename}"'

        },

    )


# ============================================================
# DELETE RESUME
# ============================================================

# @app.delete("/resumes/{resume_id}")
# async def delete_resume(

#     resume_id: str,

#     x_user_id: str = Header(...),

# ):

#     document = await get_resume_document(

#         resume_id,

#         x_user_id,

#     )


#     try:

#         minio_client.remove_object(

#             MINIO_BUCKET,

#             document["object_name"],

#         )

#     except S3Error as exc:

#         raise HTTPException(

#             status_code=500,

#             detail=(
#                 f"Unable to delete stored file: {exc}"
#             ),

#         )


#     await db.resumes.delete_one(

#         {

#             "_id":
#                 ObjectId(resume_id),

#             "user_id":
#                 x_user_id,

#         }

#     )


#     return {

#         "message":
#             "Resume deleted successfully"

#     }


# @app.delete("/resumes/{resume_id}")
# async def delete_resume(
#     resume_id: str,
#     x_user_id: str = Header(...),
# ):
#     if not ObjectId.is_valid(resume_id):
#         raise HTTPException(
#             status_code=400,
#             detail="Invalid resume ID",
#         )

#     document = await db.resumes.find_one({
#         "_id": ObjectId(resume_id),
#         "user_id": x_user_id,
#     })

#     if not document:
#         raise HTTPException(
#             status_code=404,
#             detail="Resume not found",
#         )

#     print("DELETE REQUEST:")
#     print("resume_id =", resume_id)
#     print("user_id =", x_user_id)
#     print("object_name =", document["object_name"])

#     try:
#         minio_client.remove_object(
#             MINIO_BUCKET,
#             document["object_name"],
#         )

#     except S3Error as exc:
#         raise HTTPException(
#             status_code=500,
#             detail=f"Unable to delete stored file: {exc}",
#         )

#     result = await db.resumes.delete_one({
#         "_id": ObjectId(resume_id),
#         "user_id": x_user_id,
#     })

#     print("DELETED COUNT =", result.deleted_count)

#     if result.deleted_count != 1:
#         raise HTTPException(
#             status_code=500,
#             detail="Resume database deletion failed",
#         )

#     return {
#         "message": "Resume deleted successfully",
#         "deleted_id": resume_id,
#     }


# 

@app.delete("/resumes/{resume_id}")
async def delete_resume(
    resume_id: str,
    x_user_id: str = Header(...),
):
    document = await get_resume_document(
        resume_id,
        x_user_id,
    )

    try:
        await delete_storage_file(
            document["object_name"]
        )
    except HTTPException as exc:
        print(
            "STORAGE DELETE WARNING:",
            resume_id,
            document["object_name"],
            exc.detail,
        )

        # Continue only for old/missing storage objects.
        # Do not hide authentication or other Supabase failures.
        detail = str(exc.detail).lower()

        if (
            "not found" not in detail
            and "404" not in detail
        ):
            raise

    result = await db.resumes.delete_one(
        {
            "_id": ObjectId(resume_id),
            "user_id": x_user_id,
        }
    )

    if result.deleted_count != 1:
        raise HTTPException(
            status_code=500,
            detail="Resume database deletion failed",
        )

    return {
        "message": "Resume deleted successfully",
        "deleted_id": resume_id,
    }



# ============================================================
# HEALTH
# ============================================================

@app.get("/health")
async def health():

    return {

        "status":
            "ok",

        "service":
            "resume_parser",

    }